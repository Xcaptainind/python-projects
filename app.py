import os
import pandas as pd
import random
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import Flask, request, render_template, send_file, redirect, url_for, flash
from io import BytesIO
from docx2pdf import convert
import webbrowser
import pythoncom

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Required for flash messages

# Global variables to store questions
two_mark_questions = []
mark_based_questions = []
uploaded_file_content = None

def generate_questions(df, exam_type, mid_type=None):
    global two_mark_questions, mark_based_questions
    two_mark_questions = []
    mark_based_questions = []

    # Debugging: Print the column names and a sample of the dataframe
    print("DataFrame Columns:", df.columns)
    print("DataFrame Head:\n", df.head())

    def select_questions(questions, count):
        available_count = len(questions['Question'].tolist())
        print(f"Selecting {count} questions from available {available_count} questions.")
        if available_count < count:
            return f"Insufficient questions available."  # Indicate insufficient questions
        return random.sample(questions['Question'].tolist(), count)

    if exam_type == 'internal':
        if mid_type == 'mid-1':
            # For mid-1, focus on units 1, 2, and 3.1
            units = df[df['Unit'].isin([1, 2, 3.1])]
            for unit in [1, 2]:
                unit_questions = units[units['Unit'] == unit]
                two_mark_result = select_questions(unit_questions[unit_questions['Marks'] == 2], 2)
                if isinstance(two_mark_result, str):
                    flash(two_mark_result)
                    return redirect(url_for('internal_exam'))
                two_mark_questions += two_mark_result
            unit_questions = units[units['Unit'] == 3.1]
            two_mark_result = select_questions(unit_questions[unit_questions['Marks'] == 2], 1)
            if isinstance(two_mark_result, str):
                flash(two_mark_result)
                return redirect(url_for('internal_exam'))
            two_mark_questions += two_mark_result
            for unit in [1, 2, 3.1]:
                unit_questions = units[units['Unit'] == unit]
                mark_based_result = select_questions(unit_questions[unit_questions['Marks'].isin([4, 5, 7, 10])], 2)
                if isinstance(mark_based_result, str):
                    flash(mark_based_result)
                    return redirect(url_for('internal_exam'))
                mark_based_questions += mark_based_result
        elif mid_type == 'mid-2':
            # For mid-2, focus on units 3.2, 4, and 5
            units = df[df['Unit'].isin([3.2, 4, 5])]
            for unit in [4, 5]:
                unit_questions = units[units['Unit'] == unit]
                two_mark_result = select_questions(unit_questions[unit_questions['Marks'] == 2], 2)
                if isinstance(two_mark_result, str):
                    flash(two_mark_result)
                    return redirect(url_for('internal_exam'))
                two_mark_questions += two_mark_result
            unit_questions = units[units['Unit'] == 3.2]
            two_mark_result = select_questions(unit_questions[unit_questions['Marks'] == 2], 1)
            if isinstance(two_mark_result, str):
                flash(two_mark_result)
                return redirect(url_for('internal_exam'))
            two_mark_questions += two_mark_result
            for unit in [3.2, 4, 5]:
                unit_questions = units[units['Unit'] == unit]
                mark_based_result = select_questions(unit_questions[unit_questions['Marks'].isin([4, 5, 7, 10])], 2)
                if isinstance(mark_based_result, str):
                    flash(mark_based_result)
                    return redirect(url_for('internal_exam'))
                mark_based_questions += mark_based_result
    elif exam_type == 'external':
        # External exam logic
        units = df.groupby('Unit')
        for unit, questions in units:
            if unit in [1, 2, 3.1, 3.2, 4, 5]:
                # For part A, select 2 questions from units 1, 2, 4, and 5
                if unit in [1, 2, 4, 5]:
                    two_mark_result = select_questions(questions[questions['Marks'] == 2], 2)
                    if isinstance(two_mark_result, str):
                        flash(two_mark_result)
                        return redirect(url_for('exam_type_selection'))
                    two_mark_questions += two_mark_result
                elif unit in [3.1, 3.2]:
                    two_mark_result = select_questions(questions[questions['Marks'] == 2], 1)
                    if isinstance(two_mark_result, str):
                        flash(two_mark_result)
                        return redirect(url_for('exam_type_selection'))
                    two_mark_questions += two_mark_result
        # For part B, select 2 questions from units 1, 2, 4, and 5
        for unit, questions in units:
            if unit in [1, 2, 4, 5]:
                mark_based_result = select_questions(questions[questions['Marks'].isin([4, 5, 7, 10])], 2)
                if isinstance(mark_based_result, str):
                    flash(mark_based_result)
                    return redirect(url_for('exam_type_selection'))
                mark_based_questions += mark_based_result
        # For units 3.1 and 3.2 combined, select 2 questions
        combined_3x = df[(df['Unit'] == 3.1) | (df['Unit'] == 3.2)]
        mark_based_result = select_questions(combined_3x[combined_3x['Marks'].isin([4, 5, 7, 10])], 2)
        if isinstance(mark_based_result, str):
            flash(mark_based_result)
            return redirect(url_for('exam_type_selection'))
        mark_based_questions += mark_based_result
@app.route('/', methods=['GET', 'POST'])
def upload_file():
    global uploaded_file_content
    if request.method == 'POST':
        if 'file' not in request.files:
            return 'No file part'
        file = request.files['file']
        if file.filename == '':
            return 'No selected file'
        if file:
            uploaded_file_content = file.read()
            return redirect(url_for('exam_type_selection'))
    return render_template('upload.html')

@app.route('/exam_type', methods=['GET', 'POST'])
def exam_type_selection():
    if request.method == 'POST':
        exam_type = request.form['exam_type']
        if exam_type == 'internal':
            return redirect(url_for('internal_exam'))
        elif exam_type == 'external':
            df = pd.read_excel(BytesIO(uploaded_file_content), engine='openpyxl')
            generate_questions(df, 'external')
            return redirect(url_for('preview_questions'))
    return render_template('exam_type.html')

@app.route('/internal_exam', methods=['GET', 'POST'])
def internal_exam():
    if request.method == 'POST':
        mid_type = request.form['mid_type']
        df = pd.read_excel(BytesIO(uploaded_file_content), engine='openpyxl')
        generate_questions(df, 'internal', mid_type)
        return redirect(url_for('preview_questions'))
    return render_template('internal_exam.html')

@app.route('/preview', methods=['GET', 'POST'])
def preview_questions():
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'yes':
            return redirect(url_for('input_details'))
        elif action == 'no':
            df = pd.read_excel(BytesIO(uploaded_file_content), engine='openpyxl')  # Re-read the in-memory content
            return redirect(url_for('exam_type_selection'))
        elif action == 'upload_new':
            return redirect(url_for('upload_file'))
    return render_template('preview.html', two_mark_questions=two_mark_questions, mark_based_questions=mark_based_questions)

@app.route('/input_details', methods=['GET', 'POST'])
def input_details():
    if request.method == 'POST':
        date = request.form['date']
        time = request.form['time']
        session = request.form['session']
        subject = request.form['subject']
        branch_list = request.form.getlist('branch')
        branch = ", ".join(branch_list)
        filename = generate_document(date, time, session, subject, branch)
        pdf_filename = convert_to_pdf(filename)
        file_path = os.path.join(os.getcwd(), pdf_filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return 'File not found'
    return render_template('input_details.html')

from docx.shared import Pt  # Import Pt for setting font size

from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Inches,Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_document(date, time, session, subject, branch):
    docx_filename = 'Selected_Questions.docx'

    # Delete the existing file if it exists
    if os.path.exists(docx_filename):
        os.remove(docx_filename)

    doc = Document()

    # Add logo and main heading
    logo_path = r'C:\Users\Aslam\OneDrive\Documents\programs\python\project1\templates\logo.png'  # Update with your logo path
    section = doc.sections[0]
    header = section.header

    # Specify table column widths
    table = header.add_table(rows=1, cols=2, width=Inches(6))  # Define total table width

    # Set column widths
    table.columns[0].width = Inches(2)  # First column width
    table.columns[1].width = Inches(4)  # Second column width

    # Left cell: Insert logo
    left_cell = table.cell(0, 0)
    paragraph_left = left_cell.paragraphs[0]
    run_left = paragraph_left.add_run()
    run_left.add_picture(logo_path, width=Inches(1))

    # Right cell: Insert main heading text
    right_cell = table.cell(0, 1)
    paragraph_right = right_cell.paragraphs[0]
    paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_right = paragraph_right.add_run("RV Institute of Technology")
    run_right.font.size = Pt(20)  # Adjust font size for the main heading

    # Sub-heading
    sub_heading_paragraph = doc.add_paragraph()
    sub_heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sub_heading = sub_heading_paragraph.add_run(
        "B.Tech I Semester I Mid Examinations - October 2024"
    )
    run_sub_heading.font.size = Pt(12)

    # Date, Time, Session, and Subject
    date_time_paragraph = doc.add_paragraph(f'Date: {date}                              \t\t\t\t\t Time: {time}')
    date_time_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    session_subject_paragraph = doc.add_paragraph(f'Session: {session}                        \t\t\t\t\t  Subject: {subject}')
    session_subject_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    time_marks_paragraph = doc.add_paragraph(f'Time: 2 Hours                         \t\t\t\t\t Max. Marks: 25')
    time_marks_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Part A
    doc.add_paragraph("PART - A", style='Heading1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('ANSWER ALL QUESTIONS       \t\t\t\t\t5 x 2 = 10 Marks').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    question_number = 1
    for question in two_mark_questions:
        doc.add_paragraph(f"{question_number}. {question}")
        question_number += 1

    # Part B
    doc.add_paragraph("PART - B", style='Heading1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('ANSWER ANY 5 QUESTIONS     \t\t\t\t\t  5 x 3 = 15 Marks').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    question_number = 1
    for question in mark_based_questions:
        doc.add_paragraph(f"{question_number}. {question}")
        question_number += 1

    # Save the document
    doc.save(docx_filename)
    print(f"Document saved as {docx_filename}")
    return docx_filename

def convert_to_pdf(docx_filename):
    pdf_filename = docx_filename.replace('.docx', '.pdf')

    # Delete the existing file if it exists
    if os.path.exists(pdf_filename):
        os.remove(pdf_filename)
    
    # Initialize COM library
    pythoncom.CoInitialize()

    convert(docx_filename, pdf_filename)
    print(f"Converted to PDF: {pdf_filename}")
    return pdf_filename

if __name__ == "__main__":
    port = 5000
    url = f"http://127.0.0.1:{port}/"
    webbrowser.open(url)
    app.run(debug=True, port=port)